#!/usr/bin/env python
# coding: utf-8

# In[4]:


#import ctypes
#kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)
#process_array = (ctypes.c_uint8 * 1)()
#num_processes = kernel32.GetConsoleProcessList(process_array, 1)
#if num_processes < 3: ctypes.WinDLL('user32').ShowWindow(kernel32.GetConsoleWindow(), 0)

#=============GUI========
from datetime import datetime
import time
import threading

try: import tkinter
except ImportError:
    import Tkinter as tkinter
    import ttk
else: from tkinter import ttk
from tkinter import filedialog as fd
from tkinter.filedialog import asksaveasfile
from tkinter import HORIZONTAL, BOTTOM, NW

import os
import sys
import re
import numpy as np
#==========================pdf handling
import cv2
import pdfplumber
from pdf2image import convert_from_path
from pdf2image import pdfinfo_from_path
from pdf2image.exceptions import (
    PDFInfoNotInstalledError,
    PDFPageCountError,
    PDFSyntaxError
)

#===module for industry analysis====
try:
    from googlesearch import search
except ImportError:
    print("No module named 'google' found")
#from newspaper import Article
#import requests
#from urllib.request import urlopen, Request    
from docx import Document

date = datetime.strptime('2023-11-29 23:59:59', '%Y-%m-%d %H:%M:%S')
today = datetime.now()
old_stdout = sys.stdout

class TextRedirector(object):
    def __init__(self, widget, tag="stdout"):
        self.widget = widget
        self.tag = tag
    def write(self, str):
        self.widget.insert("end", str, (self.tag,))
        self.widget.see('end')    
        
class GUI_Core():
    def __init__(self):

        def changed(value):
            if value == "Validate": self.permission()             
            elif value == "u-turn": self.start_thread2() 
            elif value == "related parties": self.related_parties()                 
            elif value == "major buyers": self.major_buyers() 
            elif value == "major suppliers": self.major_suppliers()     
            elif value == "any names": self.any_names()                  
            else: gui.text.insert("end", "under development", "red", "\n")
        
        self.root = tkinter.Tk()
        self.root.geometry("1100x500") 
        self.root.title("Vision v.1.2")
        self.root.resizable(False, False) #self.root.attributes("-toolwindow", 1)
        self.root.iconbitmap('icon.ico')
        
        self.text = tkinter.Text(self.root, height = 20, width = 132, wrap="word" )
        self.text.pack()
        sys.stdout = TextRedirector(self.text, "stdout")
        self.text.tag_configure("red", foreground="#EE0000") ##b22222
        self.text.tag_configure("blue", foreground="blue")
        self.text.tag_configure("green", foreground="#008000")
        self.text.tag_configure("black", foreground="black")
        self.text.insert("end", "Welcome to Vision v.1.2\n- Extract SLIK\n- Validate bank statements\n- Analyze transactions: u-turn, related parties, major buyers/supplier, & any specific names\n- Industry analysis\n", "blue", "\n")
     
        self.progbar = ttk.Progressbar(self.root)
        self.progbar.config(length = 1060, maximum=10, mode='indeterminate')
        self.progbar.pack()   
        #canvas untuk filename
        self.canvas1 = tkinter.Canvas(self.root, width = 100, height = 5).pack()
        self.label1 = ttk.Label(self.root)
        self.label1.pack()
        #SLIK button    
        self.canvas2 = tkinter.Canvas(self.root, width = 100, height = 5).pack()     
        self.b_startV = ttk.Button(self.root)
        self.b_startV.config(text='SLIK', command = self.permissionSLIK)
        self.b_startV.place(x = 405, y = 412, height= 50, width=70)  
        #Bank Statement button    
        self.canvas3 = tkinter.Canvas(self.root, width = 100, height = 5).pack()          
        self.menu = tkinter.StringVar(self.root)
        self.menu.set("Bank Statement")
        self.drop = tkinter.OptionMenu(self.root, self.menu, "Validate", "u-turn", "related parties","major buyers","major suppliers","any names", command=changed)
        self.drop.config(width=13)
        self.drop.place(x = 485, y = 410)       
        #Industry button    
        self.canvas2 = tkinter.Canvas(self.root, width = 100, height = 5).pack()     
        self.b_Industry = ttk.Button(self.root)
        self.b_Industry.config(text='Industry', width=15, command = self.industry)
        self.b_Industry.place(x = 620, y = 412, height= 50, width=70)        
        #supported_PDF button
        self.b_sup = ttk.Button(self.root)
        self.b_sup.config(text='F', width=1.5, command = self.supported_PDF)
        self.b_sup.place(x = 1080, y = 470)       
        #canvas untuk status result
        self.canvas3 = tkinter.Canvas(self.root, width = 100, height = 5).pack()
        self.label3 = ttk.Label(self.root)
        self.label3.place(x = 500, y = 380)               
        #link to output file0 RIGHT   
        self.link = tkinter.Label(self.root)       
        self.link.place(x = 710, y = 450) #pack(side="right")   
        #link to output file1 CENTER LOW   
        self.link1 = tkinter.Label(self.root)       
        self.link1.place(x = 497, y = 450)   
        #link to input file  LEFT  
        self.link2 = tkinter.Label(self.root)       
        self.link2.place(x = 290, y = 450) #pack(side="left")            
        #copyright label    
        self.label5 = tkinter.Label(self.root)
        self.label5.config(text='Vision v.1.2 - created in 2022 by Antonius Suwarji (antoniussuwarji@hsbc.co.id)', font=('helvetica', 8))
        self.label5.pack(side="bottom")
        

    def permissionSLIK(self): #permit pass button          
        if today < date:
            self.start_threadSLIK()
        else:
            self.password(1)    
    
    def permission(self): #permit pass button          
        if today < date:
            self.start_thread()
        else:
            self.password(2)
        
    def password(self, op): 
        def show():
            Y = re.search('(\d\d\d\d)-(\d\d)-(\d\d).*', str(today)).group(1)
            M = re.search('(\d\d\d\d)-(\d\d)-(\d\d).*', str(today)).group(2)
            D = re.search('(\d\d\d\d)-(\d\d)-(\d\d).*', str(today)).group(3)
            s = str(tryme)
            p = password.get() #get password from entry
            
            if p == s and op == 1:
                self.pasw.destroy()
                self.start_threadSLIK()            
            elif p == s and op == 2:
                self.pasw.destroy()
                self.start_thread()
                
        self.pasw = tkinter.Toplevel(self.root)
        self.pasw.geometry("200x50")
        self.pasw.title("input password")
        self.pasw.attributes("-toolwindow", 1)
        x = self.root.winfo_x()
        y = self.root.winfo_y()
        self.pasw.geometry("+%d+%d" %(x+450,y+300))         
        
        password = tkinter.StringVar(self.pasw) #Password variable
        passEntry = tkinter.Entry(self.pasw, textvariable=password, show='*')
        submit = tkinter.Button(self.pasw, text='GO',command=show)
        passEntry.pack() 
        submit.pack() 
        
#=========================================================================================================================
    def close_state(self):
        self.b_startV['state'] = 'normal'
        self.b_Industry['state'] = 'normal'
        self.drop['menu'].entryconfig(0, state='normal')
        self.drop['menu'].entryconfig(1, state='normal')
        self.drop['menu'].entryconfig(2, state='normal')
        self.drop['menu'].entryconfig(3, state='normal')
        self.drop['menu'].entryconfig(4, state='normal')
        self.drop['menu'].entryconfig(5, state='normal')            
        gui.label3.config(text='process aborted!',  foreground="red")
            
    def disable_button(self):
        self.b_startV['state'] = 'disable'
        self.b_Industry['state'] = 'disable'
        self.drop['menu'].entryconfig(0, state='disable')
        self.drop['menu'].entryconfig(1, state='disable')
        self.drop['menu'].entryconfig(2, state='disable')
        self.drop['menu'].entryconfig(3, state='disable')
        self.drop['menu'].entryconfig(4, state='disable')
        self.drop['menu'].entryconfig(5, state='disable')        
        gui.label1.config(text='')
        gui.label3.config(text='')
        gui.link.config(text='')
        gui.link1.config(text='')
        gui.link2.config(text='')        
        
    def industry(self):

        def printInput(event):
            inp = inputtxt.get("1.0", "end-1c")# get lines into string
            keywords = [i.strip() for i in inp.splitlines() if i!=""]
            #inputtxt.delete("1.0","end")
            #print(relpar_names)
            Button6.config(text = "Search", command = lambda:[self.start_thread7(keywords) ,self.top6.destroy()] )
        
        def temp_text(event):
            inputtxt.delete(1.0,"end")
            inputtxt.config(fg = 'black')
        def temp_text2(event):
            inputtxt.config(fg = 'black')
            return 'break'
            
        #make frame
        self.top6 = tkinter.Toplevel(self.root)
        self.top6.geometry("400x50")
        self.top6.title("keywords")
        self.top6.resizable(False, False) #self.root.attributes("-toolwindow", 1)
        self.top6.iconbitmap('icon.ico')
        
        x = self.root.winfo_x()
        y = self.root.winfo_y()
        self.top6.geometry("+%d+%d" %(x+350,y+300))                
        self.disable_button()
        #button Analyze        
        Button6 = tkinter.Button(self.top6, text = "Search")
        Button6.bind('<Button-1>', printInput)
        Button6.pack(side = "bottom")  
               
        #text box
        inputtxt = tkinter.Text(self.top6, undo=True, height = 20, width = 50, wrap=None)
        inputtxt.bind(printInput) #instead of: inputtxt.bind("<Return>", printInput) --> respond to Enter key
        inputtxt.insert(1.0, "type keywords..")
        inputtxt.config(fg = 'grey')
        inputtxt.bind("<FocusIn>", temp_text)   
        inputtxt.bind("<Return>", temp_text2) 
        inputtxt.pack(side = "left")
        
        def on_closing():
            self.close_state()
            self.top6.destroy()        
        self.top6.protocol("WM_DELETE_WINDOW", on_closing)        
        
#=========================================================================================================================

    def any_names(self):

        def printInput(event):
            inp = inputtxt.get("1.0", "end-1c")# get lines into string
            any_names = [i.strip() for i in inp.splitlines() if i!=""]
            #inputtxt.delete("1.0","end")
            #print(relpar_names)
            Button5.config(text = "Analyze", command = lambda:[self.start_thread6(any_names) ,self.top5.destroy()] )
            
        def temp_text(event):
            inputtxt.delete(1.0,"end")
            inputtxt.config(fg = 'black')
            
        #make frame
        self.top5 = tkinter.Toplevel(self.root)
        self.top5.geometry("400x400")
        self.top5.title("input any names ")
        self.top5.resizable(False, False) #self.root.attributes("-toolwindow", 1)
        self.top5.iconbitmap('icon.ico')
        
        x = self.root.winfo_x()
        y = self.root.winfo_y()
        self.top5.geometry("+%d+%d" %(x+350,y+0))        
        self.disable_button()
        #button Analyse      
        Button5 = tkinter.Button(self.top5, text = "Analyze")
        Button5.bind('<Button-1>', printInput)
        Button5.pack(side = "bottom")  
               
        #text box
        scrollbar = tkinter.Scrollbar(self.top5)
        inputtxt = tkinter.Text(self.top5, undo=True, height = 50,width = 50, wrap=None,  yscrollcommand=scrollbar.set)
        inputtxt.bind(printInput) #instead of: inputtxt.bind("<Return>", printInput) --> respond to Enter key
        inputtxt.insert(1.0, "type & press Enter or copy & paste..")
        inputtxt.config(fg = 'grey')
        inputtxt.bind("<FocusIn>", temp_text)
        scrollbar.config(command=inputtxt.yview)
        scrollbar.pack(side="right", fill="both")        
        inputtxt.pack(side = "left")
        
        def on_closing():
            self.close_state()
            self.top5.destroy()        
        self.top5.protocol("WM_DELETE_WINDOW", on_closing)            
        
        
#=========================================================================================================================

    def major_suppliers(self):

        def printInput(event):
            inp = inputtxt.get("1.0", "end-1c")# get lines into string
            majsup_names = [i.strip() for i in inp.splitlines() if i!=""]
            #inputtxt.delete("1.0","end")
            #print(relpar_names)
            Button4.config(text = "Analyze", command = lambda:[self.start_thread5(majsup_names) ,self.top4.destroy()] )

        def temp_text(event):
            inputtxt.delete(1.0,"end")
            inputtxt.config(fg = 'black')            
            
        #make frame
        self.top4 = tkinter.Toplevel(self.root)
        self.top4.geometry("400x400")
        self.top4.title("input major suppliers")
        self.top4.resizable(False, False) #self.root.attributes("-toolwindow", 1)
        self.top4.iconbitmap('icon.ico')
        
        x = self.root.winfo_x()
        y = self.root.winfo_y()
        self.top4.geometry("+%d+%d" %(x+350,y+0))        
        self.disable_button()
        #button Analyze        
        Button4 = tkinter.Button(self.top4, text = "Analyze")
        Button4.bind('<Button-1>', printInput)
        Button4.pack(side = "bottom")  
               
        #text box
        scrollbar = tkinter.Scrollbar(self.top4)
        inputtxt = tkinter.Text(self.top4, undo=True, height = 50,width = 50, wrap=None,  yscrollcommand=scrollbar.set)
        inputtxt.bind(printInput) #instead of: inputtxt.bind("<Return>", printInput) --> respond to Enter key
        inputtxt.insert(1.0, "type & press Enter or copy & paste..")
        inputtxt.config(fg = 'grey')
        inputtxt.bind("<FocusIn>", temp_text)
        scrollbar.config(command=inputtxt.yview)
        scrollbar.pack(side="right", fill="both")        
        inputtxt.pack(side = "left")
        
        def on_closing():
            self.close_state()
            self.top4.destroy()        
        self.top4.protocol("WM_DELETE_WINDOW", on_closing)            
        
        
#=========================================================================================================================

    def major_buyers(self):

        def printInput(event):
            inp = inputtxt.get("1.0", "end-1c")# get lines into string
            majbuy_names = [i.strip() for i in inp.splitlines() if i!=""]
            #inputtxt.delete("1.0","end")
            #print(relpar_names)
            Button3.config(text = "Analyze", command = lambda:[self.start_thread4(majbuy_names) ,self.top3.destroy()] )

        def temp_text(event):
            inputtxt.delete(1.0,"end")
            inputtxt.config(fg = 'black')            
            
        #make frame
        self.top3 = tkinter.Toplevel(self.root)
        self.top3.geometry("400x400")
        self.top3.title("input major buyers")
        self.top3.resizable(False, False) #self.root.attributes("-toolwindow", 1)
        self.top3.iconbitmap('icon.ico')
        
        x = self.root.winfo_x()
        y = self.root.winfo_y()
        self.top3.geometry("+%d+%d" %(x+350,y+0))        
        self.disable_button()
        #button Analyze       
        Button3 = tkinter.Button(self.top3, text = "Analyze")
        Button3.bind('<Button-1>', printInput)
        Button3.pack(side = "bottom")  
               
        #text box
        scrollbar = tkinter.Scrollbar(self.top3)
        inputtxt = tkinter.Text(self.top3, undo=True, height = 50,width = 50, wrap=None,  yscrollcommand=scrollbar.set)
        inputtxt.bind(printInput) #instead of: inputtxt.bind("<Return>", printInput) --> respond to Enter key
        inputtxt.insert(1.0, "type & press Enter or copy & paste..")
        inputtxt.config(fg = 'grey')
        inputtxt.bind("<FocusIn>", temp_text)
        scrollbar.config(command=inputtxt.yview)
        scrollbar.pack(side="right", fill="both")        
        inputtxt.pack(side = "left")
        
        def on_closing():
            self.close_state()
            self.top3.destroy()        
        self.top3.protocol("WM_DELETE_WINDOW", on_closing)            
    
    
#=========================================================================================================================
    def related_parties(self):

        def printInput(event):
            inp = inputtxt.get("1.0", "end-1c")# get lines into string
            relpar_names = [i.strip() for i in inp.splitlines() if i!=""]
            #inputtxt.delete("1.0","end")
            #print(relpar_names)
            Button2.config(text = "Analyze", command = lambda:[self.start_thread3(relpar_names) ,self.top2.destroy()] )

        def temp_text(event):
            inputtxt.delete(1.0,"end")
            inputtxt.config(fg = 'black')
            
        #make frame
        self.top2 = tkinter.Toplevel(self.root)
        self.top2.geometry("400x400")
        self.top2.title("input related parties")
        self.top2.resizable(False, False) #self.root.attributes("-toolwindow", 1)
        self.top2.iconbitmap('icon.ico')
        
        x = self.root.winfo_x()
        y = self.root.winfo_y()
        self.top2.geometry("+%d+%d" %(x+350,y+0))        
        self.disable_button()
        #button Analyze        
        Button2 = tkinter.Button(self.top2, text = "Analyze")
        Button2.bind('<Button-1>', printInput)
        Button2.pack(side = "bottom")  
               
        #text box
        scrollbar = tkinter.Scrollbar(self.top2)
        inputtxt = tkinter.Text(self.top2, undo=True, height = 50,width = 50, wrap=None,  yscrollcommand=scrollbar.set)
        inputtxt.bind(printInput) #instead of: inputtxt.bind("<Return>", printInput) --> respond to Enter key
        inputtxt.insert(1.0, "type & press Enter or copy & paste..")
        inputtxt.config(fg = 'grey')
        inputtxt.bind("<FocusIn>", temp_text)
        scrollbar.config(command=inputtxt.yview)
        scrollbar.pack(side="right", fill="both")        
        inputtxt.pack(side = "left")
        
        def on_closing():
            self.close_state()
            self.top2.destroy()        
        self.top2.protocol("WM_DELETE_WINDOW", on_closing)
        
#=========================================================================================================================        
        
    def supported_PDF(self):
        self.top1 = tkinter.Toplevel(self.root)
        self.top1.geometry("900x500")
        self.top1.title("Guidance")
        self.top1.resizable(False, False) #self.root.attributes("-toolwindow", 1)
        self.top1.iconbitmap('icon.ico')
        
        def nex_img(i):   # takes the current scale position as an argument
            # delete previous image
            canvas.delete('image')
            # create next image
            canvas.create_image(0, 0, anchor=NW, image=listimg[int(i)-1], tags='image')            
            
        if getattr(sys, 'frozen', False):
            imageA = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/01-SLIK1.png"))
            imageB = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/02-SLIK2.png"))
            imageC = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/03-SLIK3.png"))
            imageD = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/04-SLIK4.png"))
            imageE = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/05-BS1.png"))
            imageF = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/06-BS2.png"))
            imageG = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/07-BS3.png"))
            imageH = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/08-BS4.png"))
            imageI = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/09-BS5.png"))
            imageJ = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/10-BS6.png"))
            imageK = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/11-BS7.png"))
            imageL = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/12-BS8.png"))
            imageM = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/13-BS9.png"))
            imageN = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/14-BS10.png"))
            imageO = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/15-BS11.png"))
            imageP = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/16-BS12.png"))
            imageQ = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/17-BS13.png"))
            imageR = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/18-BS14.png"))
            imageS = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/19-BS15.png"))
            imageT = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/20-IND1.png"))
            imageU = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/21-IND2.png"))
            imageV = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/22-IND3.png"))
            imageW = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/23-IND4.png"))
            image1 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BCA1.png"))
            image2 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BCA2.png"))
            image3 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BCA2-other.png"))
            image4 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BCA1-scan.png"))
            image5 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BRI1.png"))
            image6 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BRI2.png"))
            image7 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BRI1-scan.png"))
            image8 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/MANDIRI1.png"))
            image9 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/MANDIRI2.png"))
            image10 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/MANDIRI3.png"))
            image11 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/MANDIRI4-scan.png"))
            image12 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/MANDIRI5-scan.png"))
            image13 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BNI1.png"))
            image14 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BNI2.png"))
            image15 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BNI2-scan.png"))
            image16 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/BNI3-scan.png"))
            image17 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/PERMATA1.png"))
            image18 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/CIMB1.png"))
            image19 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/CIMB2.png"))
            image20 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/OCBC1.png"))
            image21 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/OCBC2.png"))
            image22 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/DBS1.png"))
            image23 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/QNB1.png"))
            image24 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/UOB1.png"))
            image25 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/DANAMON1.png"))
            image26 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/DANAMON-scan.png"))
            image27 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/HSBC1.png"))
            image28 = tkinter.PhotoImage(file=os.path.join(sys._MEIPASS, "doc/SLIK.png"))            
        else:
            imageA = tkinter.PhotoImage(file= "doc/01-SLIK1.png")
            imageB = tkinter.PhotoImage(file= "doc/02-SLIK2.png")
            imageC = tkinter.PhotoImage(file= "doc/03-SLIK3.png")
            imageD = tkinter.PhotoImage(file= "doc/04-SLIK4.png")
            imageE = tkinter.PhotoImage(file= "doc/05-BS1.png")
            imageF = tkinter.PhotoImage(file= "doc/06-BS2.png")
            imageG = tkinter.PhotoImage(file= "doc/07-BS3.png")
            imageH = tkinter.PhotoImage(file= "doc/08-BS4.png")
            imageI = tkinter.PhotoImage(file= "doc/09-BS5.png")
            imageJ = tkinter.PhotoImage(file= "doc/10-BS6.png")
            imageK = tkinter.PhotoImage(file= "doc/11-BS7.png")
            imageL = tkinter.PhotoImage(file= "doc/12-BS8.png")
            imageM = tkinter.PhotoImage(file= "doc/13-BS9.png")
            imageN = tkinter.PhotoImage(file= "doc/14-BS10.png")
            imageO = tkinter.PhotoImage(file= "doc/15-BS11.png")
            imageP = tkinter.PhotoImage(file= "doc/16-BS12.png")
            imageQ = tkinter.PhotoImage(file= "doc/17-BS13.png")
            imageR = tkinter.PhotoImage(file= "doc/18-BS14.png")
            imageS = tkinter.PhotoImage(file= "doc/19-BS15.png")
            imageT = tkinter.PhotoImage(file= "doc/20-IND1.png")
            imageU = tkinter.PhotoImage(file= "doc/21-IND2.png")
            imageV = tkinter.PhotoImage(file= "doc/22-IND3.png")
            imageW = tkinter.PhotoImage(file= "doc/23-IND4.png")           
            image1 = tkinter.PhotoImage(file= "doc/BCA1.png")
            image2 = tkinter.PhotoImage(file = "doc/BCA2.png")
            image3 = tkinter.PhotoImage(file = "doc/BCA2-other.png")
            image4 = tkinter.PhotoImage(file = "doc/BCA1-scan.png")
            image5 = tkinter.PhotoImage(file = "doc/BRI1.png")
            image6 = tkinter.PhotoImage(file = "doc/BRI2.png")
            image7 = tkinter.PhotoImage(file = "doc/BRI1-scan.png")
            image8 = tkinter.PhotoImage(file = "doc/MANDIRI1.png")
            image9 = tkinter.PhotoImage(file = "doc/MANDIRI2.png")
            image10 = tkinter.PhotoImage(file = "doc/MANDIRI3.png")
            image11 = tkinter.PhotoImage(file = "doc/MANDIRI4-scan.png")
            image12 = tkinter.PhotoImage(file = "doc/MANDIRI5-scan.png")
            image13 = tkinter.PhotoImage(file = "doc/BNI1.png")
            image14 = tkinter.PhotoImage(file = "doc/BNI2.png") 
            image15 = tkinter.PhotoImage(file = "doc/BNI2-scan.png") 
            image16 = tkinter.PhotoImage(file = "doc/BNI3-scan.png") 
            image17 = tkinter.PhotoImage(file = "doc/PERMATA1.png")
            image18 = tkinter.PhotoImage(file = "doc/CIMB1.png")
            image19 = tkinter.PhotoImage(file = "doc/CIMB2.png")
            image20 = tkinter.PhotoImage(file = "doc/OCBC1.png")
            image21 = tkinter.PhotoImage(file = "doc/OCBC2.png")
            image22 = tkinter.PhotoImage(file = "doc/DBS1.png")
            image23 = tkinter.PhotoImage(file = "doc/QNB1.png")
            image24 = tkinter.PhotoImage(file = "doc/UOB1.png")
            image25 = tkinter.PhotoImage(file = "doc/DANAMON1.png")
            image26 = tkinter.PhotoImage(file = "doc/DANAMON-scan.png")
            image27 = tkinter.PhotoImage(file = "doc/HSBC1.png")
            image28 = tkinter.PhotoImage(file = "doc/SLIK.png")        
        listimg = [imageA, imageB, imageC, imageD, imageE, imageF, imageG, imageH, imageI, imageJ, imageK, imageL, imageM, imageN, imageO, imageP, imageQ, imageR, imageS, imageT, imageU, imageV, imageW, image1, image2, image3, image4, image5, image6, image7, image8, image9, image10, image11, image12, image13, image14, image15, image16, image17, image18, image19, image20, image21, image22, image23, image24, image25, image26, image27, image28]
        scale = tkinter.Scale(master=self.top1, orient=HORIZONTAL, length = 200, from_=1, to=len(listimg), resolution=1, showvalue=False, command=nex_img)
        scale.pack(side=BOTTOM)        
        canvas = tkinter.Canvas(self.top1, width=875, height=550)
        canvas.pack(fill="both", expand=True)    
        nex_img(1)
        
#=========================================================================================================================
    def thread_state(self):
        self.b_startV['state'] = 'disable'
        self.b_Industry['state'] = 'disable'
        self.drop['menu'].entryconfig(0, state='disable')
        self.drop['menu'].entryconfig(1, state='disable')
        self.drop['menu'].entryconfig(2, state='disable')
        self.drop['menu'].entryconfig(3, state='disable')
        self.drop['menu'].entryconfig(4, state='disable')
        self.drop['menu'].entryconfig(5, state='disable')          
        
    def start_threadSLIK(self): #SLIK button     
        self.thread_state()
        self.secondary_thread = threading.Thread(target=arbitrarySLIK)
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)])  
        
    def start_thread(self): #Validate button     
        self.thread_state()      
        self.secondary_thread = threading.Thread(target=arbitrary)
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)])   

    def start_thread2(self): #Analyze u-turn     
        self.thread_state()      
        self.secondary_thread = threading.Thread(target=arbitrary2)
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)]) 
        
    def start_thread3(self, relpar_names): #Analyze related parties
        self.thread_state()      
        self.secondary_thread = threading.Thread(target=lambda: arbitrary3(relpar_names))
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)]) 
        
    def start_thread4(self, majbuy_names): #Analyze major buyers
        self.thread_state()      
        self.secondary_thread = threading.Thread(target=lambda: arbitrary4(majbuy_names))
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)])
        
    def start_thread5(self, majsup_names): #Analyze major suppliers
        self.thread_state()       
        self.secondary_thread = threading.Thread(target=lambda: arbitrary5(majsup_names))
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)])        
        
    def start_thread6(self, any_names): #Analyze any names
        self.thread_state()       
        self.secondary_thread = threading.Thread(target=lambda: arbitrary6(any_names))
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)])           

    def start_thread7(self, keywords): #Analyze any names
        self.thread_state()        
        self.secondary_thread = threading.Thread(target=lambda: arbitrary7(keywords))
        self.secondary_thread.start()  
        self.root.after(20, lambda:[check_thread(0)])  
        
def check_thread(x):
    if gui.secondary_thread.is_alive():
        x+=1
        gui.root.after(20, lambda:[check_thread(x)])
    else:
        gui.progbar.stop()
        gui.b_startV['state'] = 'normal'
        gui.b_Industry['state'] = 'normal'
        gui.drop['menu'].entryconfig(0, state='normal')
        gui.drop['menu'].entryconfig(1, state='normal')
        gui.drop['menu'].entryconfig(2, state='normal')    
        gui.drop['menu'].entryconfig(3, state='normal')
        gui.drop['menu'].entryconfig(4, state='normal')
        gui.drop['menu'].entryconfig(5, state='normal')        
        print("")            
    return x
   
    
#==========================================================   

def func_aSLIK():
    
    #clear all status
    gui.label1.config(text='')
    gui.label3.config(text='')
    gui.link.config(text='')
    gui.link1.config(text='')
    gui.link2.config(text='')
    
    #open dialog window and read filepath
    filepath= fd.askopenfilename(title='Open PDF file', filetypes=(('pdf files', '*.pdf'),))
    onlyfilename = os.path.basename(filepath) 
    onlypath = os.path.dirname(filepath) #filepath=onlypath + "/" + onlyfilename
    gui.label1.config(text='file selected : '+ onlyfilename, foreground="blue")
    print("")
    print("source: ",filepath)
    #start progbar
    gui.progbar.start()
    
    #try open pdf with pdfplumber
    try:              
        with pdfplumber.open(filepath) as pdf:
            try:
                newpath = onlypath + "/" + "temp2410"
                if not os.path.exists(newpath):
                    os.mkdir(newpath)
                #clean directory temp2410
                file_xls = [f for f in os.listdir(newpath) if f.endswith('.xlsx')]
                for f in file_xls: os.remove(os.path.join(newpath, f))                    
                file_jpg = [f for f in os.listdir(newpath) if f.endswith('.jpg')]
                for f in file_jpg: os.remove(os.path.join(newpath, f))
                        
            except PermissionError:
                gui.text.insert("end", "please close excel or jpg file output", "red", "\n")
                gui.label3.config(text='process aborted!',  foreground="red")
                print("")
                sys.exit(1)        
            
            #find page length
            info = pdfinfo_from_path(filepath, poppler_path=r'C:\Program Files\Vision\poppler-0.68.0\bin')
            lenPages = info["Pages"] 
            #extract page1 using pdfplumber
            page0 = pdf.pages[0]
            text = page0.extract_text() #text in page1
            
        if text == None or text == "":              
            print("")
            gui.text.insert("end", "unable to recognize pdf template..", "red", "\n")             
            
        elif text != None or text != "":            
            if re.findall(r"Sistem Layanan Informasi Keuangan|Informasi Debitur|Data Pokok Debitur|Penyajian informasi debitur pada Sistem Layanan Informasi Keuangan|Kredit/Pembiayaan|Tanggal Akad Awal|Pelapor", text):
                #SLIK
                print("SLIK file detected...")
                print("processing pdfplumber...")
                from SLIK import main_SLIK
                result = main_SLIK(filepath, lenPages)
                if result == None:
                    gui.label3.config(text='extraction failed!',  foreground="red")                           
                else:
                    gui.label3.config(text='extraction done !',  foreground="green") #ganti label
                    #show link to output
                    gui.link.bind("<Button-1>", lambda event, : os.startfile(result, 'open') )                
                    gui.link.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))     
                    gui.link2.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
                    gui.link2.config(text='open file input', fg="green", cursor="hand2", font=('Calibri',11,'underline'))                    
                    
            #======================none of above======================
            else:
                print("")
                gui.text.insert("end", "unable to recognize pdf template..", "red", "\n") 
                print("") 
                gui.label3.config(text='extraction failed!',  foreground="red")      
    
    except FileNotFoundError:
        gui.text.insert("end", "No file selected", "red", "\n")
        print("")
        gui.label3.config(text='process aborted!',  foreground="red")
        pass                   
    else:
        pass     
    
    time.sleep(0.) 
            
            
def func_a():
          
    def show(result):
        if re.findall('1. Program ', result[0]):
            gui.text.insert("end", result[0], "green")       
        if re.findall('2. Balance ', result[1]):
            gui.text.insert("end", result[1], "green")
        elif re.findall('#  WARNING', result[1]):
            gui.text.insert("end", result[1], "red")
        elif re.findall('2. NOTE:', result[1]):
            gui.text.insert("end", result[1], "blue")                    
        if re.findall('3. Total', result[2]):
            gui.text.insert("end", result[2], "green")
        elif re.findall('3. WARNING', result[2]):
            gui.text.insert("end", result[2], "red")
        elif re.findall('3. NOTE', result[2]):
            gui.text.insert("end", result[2], "blue") 
        if re.findall('4. Total', result[3]):
            gui.text.insert("end", result[3], "green")
        elif re.findall('4. WARNING', result[3]):
            gui.text.insert("end", result[3], "red")
        elif re.findall('4. NOTE|   # NOTE', result[3]):
            gui.text.insert("end", result[3], "blue")          
      
    def label(result, fileresult):   
        if not re.findall('#  WARNING', result):
            gui.label3.config(text='extraction done!', foreground="green")            

    def look(result):  
        if not re.findall('#  WARNING', result[1]):
            gui.link.bind("<Button-1>", lambda event, : os.startfile(result[5], 'open') )                
            gui.link.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))       
            gui.link2.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
            gui.link2.config(text='open file input', fg="green", cursor="hand2", font=('Calibri',11,'underline'))             
    
    def fail():
        gui.text.insert("end", "unable to complete process..", "red", "\n")
        gui.label3.config(text='extraction incomplete!',  foreground="red")
        pass  
    
    def fail_recognize():
        gui.text.insert("end", "unable to complete process..", "red", "\n")
        gui.label3.config(text='extraction incomplete!',  foreground="red")
        pass          
    
    #=================================================start/body================================================
    #clear all status
    gui.label1.config(text='')
    gui.label3.config(text='')
    gui.link.config(text='')
    gui.link1.config(text='')
    gui.link2.config(text='')
    
    #open dialog window and read filepath
    filepath= fd.askopenfilename(title='Open PDF file', filetypes=(('pdf files', '*.pdf'),))
    onlyfilename = os.path.basename(filepath) 
    onlypath = os.path.dirname(filepath) #filepath=onlypath + "/" + onlyfilename
    gui.label1.config(text='file selected : '+ onlyfilename, foreground="blue")
    print("")
    print("source: ",filepath)
    #start progbar
    gui.progbar.start()
    
    #try open pdf with pdfplumber
    try: 
        with pdfplumber.open(filepath) as pdf:
            try:
                newpath = onlypath + "/" + "temp2410"
                if not os.path.exists(newpath):
                    os.mkdir(newpath)
                #clean directory temp2410
                file_xls = [f for f in os.listdir(newpath) if f.endswith('.xlsx')]
                for f in file_xls: os.remove(os.path.join(newpath, f))                    
                file_jpg = [f for f in os.listdir(newpath) if f.endswith('.jpg')]
                for f in file_jpg: os.remove(os.path.join(newpath, f))
                        
            except PermissionError:
                gui.text.insert("end", "please close excel or jpg file output", "red", "\n")
                gui.label3.config(text='process aborted!',  foreground="red")
                print("")
                sys.exit(1)        
            
            #find page length
            info = pdfinfo_from_path(filepath, poppler_path=r'C:\Program Files\Vision\poppler-0.68.0\bin')
            lenPages = info["Pages"] 
            #extract page1 using pdfplumber
            page0 = pdf.pages[0]
            text = page0.extract_text() #text in page1
            page_end1 = pdf.pages[lenPages-1]
            text_end = page_end1.extract_text() #text in last page
            #accomodate page2 and last page-1
            if lenPages>1 and text != None:
                page1 = pdf.pages[1]
                text = text + page1.extract_text() #text in page1+ page2
                page_end2 = pdf.pages[lenPages-2]  
                text_end = text_end + page_end2.extract_text()  #text in last page + last page-1
     
        #==============tesseract, if pdfplumber failed
        if text == None or text =="":    
            print("detecting pdf using tesseract...")   
            #convert pdf to jpg (page1 only)
            for i in range(1,2):  
                pages = convert_from_path(filepath, dpi=400, first_page=i, last_page = i, poppler_path=r'C:\Program Files\Vision\poppler-0.68.0\bin') 
                for page in pages: 
                    image = newpath + "/" + "page-0" + str(i) + ".jpg" 
                    print("convert pdf to image : ", image)
                    page.save(image, "JPEG")  
                    #crop page01.jpg to determine suitable model
                    image = cv2.imread(newpath + "/" + "page-0" + str(i) + ".jpg")
                    y=2000
                    h=500
                    x=0
                    w=10000
                    crop_img = image[y:y+h, x:x+w]
                    cv2.imwrite(newpath + "/" + "page-0" + str(i) + ".jpg", crop_img)
                    
            #read crop image page-01 to select model
            file = newpath + "/" + "page-01.jpg"
            print("selecting model to process :", file)
            image = cv2.imread(file) # reading image using opencv              
            
            #pilih threshold
            arranged_text =[]
            conf_val = []
            for s in range (1,4): #there are 3 model (1,2,3)
                from tesseract import Extract
                thresholds_image = Extract.pre_processing(image, newpath, s)
                model = 0
                parsed_data = Extract.parse_text(thresholds_image, model)
                mod = "{:.2f}".format(parsed_data[1])
                print("model-"+ str(s) +" score : ", mod)
                conf_val.append(parsed_data[1])
                
            #select model 
            model = int(conf_val.index(max(conf_val))) + 1  
            if abs(conf_val[0]-conf_val[model-1])<15: model = 1
            print("model selected: ", model)
              
            #identifying image
            for i in range(1,2):  
                pages = convert_from_path(filepath, dpi=400, first_page=i, last_page = i, poppler_path=r'C:\Program Files\Vision\poppler-0.68.0\bin') 
                for page in pages: 
                    image = newpath + "/" + "page-0" + str(i) + ".jpg" 
                    print("convert pdf to image : ", image)
                    page.save(image, "JPEG")              
            #read image page-01 to be identified
            file = newpath + "/" + "page-01.jpg"
            image = cv2.imread(file) # reading image using opencv 
            print("identifying image :", file)
            from tesseract import Extract
            thresholds_image = Extract.pre_processing(image, newpath, model)
            parsed_data = Extract.parse_text(thresholds_image, model)
            arranged_text_init = Extract.format_text(parsed_data[0])
            arranged_text = arranged_text + arranged_text_init
            
            try:
                line=[] 
                for l in range(len(arranged_text)):
                    #line = ' '.join(arranged_text[l]) --> line per line
                    line = np.append(line, ' '.join(arranged_text[l])) #---> combined all line

                text0 = '\n'.join(map(str, line)) 
                #print(text0)
                #==================select template - tesseract===============
                found_BCA1 = len(re.findall(r"REKENING\s*TAHAPAN|Apabila nasabah|CATATAN|Laporan Mutasi|BCA berhak|CBG\s*MUTASI\s*SALDO", text0))
                found_BRI1 = len(re.findall(r"LAPORAN TRANSAKSI| BANK BRI|Dengan Setulus Hati|(((0[1-9]|1\d|2\d|3[01])/(0?[1-9]|1[0-2])/\d{2})\s+(\d\d:\d\d:\d\d)\s+(.*)\s+(0.00).*\s+(\d+\D+(\d+,|\D+,).*\d\d|\D+(\d+,|\D+,).*\d\d|(\d+,|\D+,).*\d\d|\d{1,3}\.\d\d)K.*)|(((0[1-9]|1\d|2\d|3[01])/(0?[1-9]|1[0-2])/\d{2})\s+(\d\d:\d\d:\d\d)\s+(.*)\s+(\d+\D+(\d+,|\D+,).*\d\d|\D+(\d+,|\D+,).*\d\d|(\d+,|\D+,).*\d\d|\d{1,3}\.\d\d)D.*)", text0))
                found_MANDIRI = len(re.findall(r"MANDIR.*GIRO|Saldo.*Pemindahan|(^\w)*((\w{1,2})(/|7|/7|f|://)(\d\d|,)):*—*\.*\?*,*-*\s+(\w{1,2})(/|7|/7|f|/f|/£|://)(\d\d)\s+(.*)\s+(\D+(\d+,|\D+,).*\d\d|(\d+,|\D+,).*\d\d|\d{1,3}\.\d\d)\s+(D|D-|.D)\s+(\D+(\d+,|\D+,).*\d\d|(\d+,|\D+,).*\d\d|\d{1,3}\.\d\d)", text0)) 
                found_MANDIRI2 = len(re.findall(r"KMK.*UMUM|(^\w)*((\w{1,2})(/|7|/7|f|://)(\d\d|,)):*—*\.*\?*,*-*\s+(\w{1,2})(/|7|/7|f|/f|/£|://)(\d\d)\s+(.*)\s+(\D+(\d+,|\D+,).*\d\d|(\d+,|\D+,).*\d\d|\d{1,3}\.\d\d)\s+(K|K-|.K)\s+", text0))
                found_BNI2 = len(re.findall(r"Posting Date Effective Date Transaction|((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+\d{2}\.\d{2}\.\d{2}\s+((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+\d{2}\.\d{2}\.\d{2}\s+", text0))
                found_BNI3 = len(re.findall(r"Melayani\s*Negeri|Kebanggaan\s*Bangsa|Uraian\s*Mutasi\s*Saldo|Apabila\s*terdapat|((0?[1-9]|1\d|2\d|3[01])/(0?[1-9]|1[0-2])/\d{4})\s+((0?[1-9]|1\d|2\d|3[01])/(0?[1-9]|1[0-2])/\d{4})\s+\d{7}\s+.*\s+(\d+\D+(\d+,|\D+,).*\d\d|\D+(\d+,|\D+,).*\d\d|(\d+,|\D+,).*\d\d|\d{1,3}\.\d\d)\s*D\s+(\d+\D+(\d+,|\D+,).*\d\d|\D+(\d+,|\D+,).*\d\d|(\d+,|\D+,).*\d\d|\d{1,3}\.\d\d)", text0))
                found_PERMATA2 = len(re.findall(r"PermataBank|PermataTel|Penting!\s+Efektif", text0))
                found_CIMB1 = len(re.findall(r"Consolidated Account Statement|No. Rekening Nama Product Mata Uang Saldo|Account Number Product Name Currency Balance |SAVING SUMMARY|of PT. Bank CIMB Niaga|Laporan transaksi ini akan dikirim|your account by PT. Bank CIMB Niaga", text0))
                found_OCBC1 = len(re.findall(r"Account\s+No\s+:.*Opening\s+Balance\s+:.*|Account\s+Name.*Closing\s+Balance|Total\s+Debit.*Ledger\s+Balance|Total\s+Credit\s+.*Available\s+Balance|Value\s+Date\s+Reference\s+No.\s+Cheque\s+No.\s+Description|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+\d{10,}\s+.*\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)", text0))
                found_DBS1 = len(re.findall(r"Opening Balance.*Earmark Amount|Ledger Balance.*Overdraft Limit|Date Value Date Transaction Details  Debit Credit Running Balance|((0[1-9]|1\d|2\d|3[01])-\w{3}-\d{4})\s+((0[1-9]|1\d|2\d|3[01])-\w{3}-\d{4})\s+", text0))
                found_QNB1 = len(re.findall(r"No. Post Date Effective Date Transaction Name Debit Credit Balance Information|layanan CIB BANK QNB INDONESIA|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+\d{2,5}\s+.*\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)\s+\w+", text0))
                found_UOB1 = len(re.findall(r"CompanyID:.*AccountNumber:|AccountName:.*LedgerBalance:|AccountType:.*AvailableBalance:|AccountCurrency:.*TotalFloat:|AccountBranch:.*OverdraftFacility:|AccountNature.*Earmark|Primary.*AllocatedAmount|Movement Details - From:|Statement\s+ValueDate\s+Transaction\s+Description\s+Deposit\s+Withdrawal\s+Balance|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})", text0))
                found_DANAMON1 = len(re.findall(r"Cabang :.*BANK DANAMON|RINCIAN.*TRANSAKSI|Tgl\s+Keterangan\s+Reff\s+Debit\s+Kredit\s+Saldo|Trans.*Valuta|SALDO\s+BULAN\s+LALU|((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))\s+((\d\d)/(0[1-9]|1[0-2]))\s+.*\s+(\d|\d{12})\s+((\d+,|\D+,).*|(\d+,|\D+,).*|\d{1,3}\.\d\d)\s{,4}\s+((\d+,|\D+,).*|(\d+,|\D+,).*|\d{1,3}\.\d\d)|suku\s+bunga\s+baru\s+menjadi|PT Bank Danamon Indonesia|danamoniline.com|1-500-0990|Syarat\s+&\s+Ketentuan\s+berlaku\s+hubungi|Hello\s+Danamon|bdi.co.id", text0))
                found_HSBC1 = len(re.findall(r"Laporan Rekening$|Account Statement$|Rincian Rekening$|Account Details$|Nomor Urut Laporan|Stmt Sequence Number|Tanggal Rincian Penarikan Setoran Saldo|Date Details Withdrawals Deposits Balance|(0[1-9]|1\d|2\d|3[01])\w{3}\d{4}\s+(.*)\s+(-(\d+,).*|(\d+,).*|-\d+|\d+)|Bank HSBC Indonesia|Kantor Pusat - World Trade Center|HSBC Premier|HSBC Advance|Layanan Nasabah\s*:\s*www.hsbc.co.id|Diterbitkan oleh PT Bank HSBC Indonesia|6221 2551 4722|6221 2552 6603|1 500 700|1 500 808", text0))
                
                select_list = [found_BCA1, found_BRI1, found_MANDIRI, found_MANDIRI2, found_BNI2, found_BNI3, found_PERMATA2, found_CIMB1, found_OCBC1, found_DBS1, found_QNB1, found_UOB1, found_DANAMON1, found_HSBC1]
                print(select_list)
                select = int(select_list.index(max(select_list))) + 100
                
                #==================execute tesseract==========================
                if sum(select_list)!=0 and max(select_list) >= 2 and select == 100:                        
                    try:
                        print("BCA file detected...")
                        from BCA import main_BCA
                        result = main_BCA(filepath, select, model)
                        show(result)                           
                        label(result[1], result[5]) 
                        look(result)                             
                    except: fail()
                        
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 101:  
                    try:
                        print("BRI1 file detected...")
                        from BRI import main_BRI
                        result = main_BRI(filepath, select, model)
                        show(result)
                        label(result[1], result[5]) 
                        look(result)                                                      
                    except: fail() 
                                    
                elif sum(select_list)!=0 and max(select_list) >= 2 and select == 102:  
                    try:
                        print("MANDIRI file detected...")  
                        from MANDIRI import main_MANDIRI
                        result = main_MANDIRI(filepath, select, model)
                        show(result)
                        label(result[1], result[5]) 
                        look(result)                                                       
                    except: fail()

                elif sum(select_list)!=0 and max(select_list) >= 2 and select == 103:  
                    try:
                        print("MANDIRI2 file detected...")  
                        from MANDIRI import main_MANDIRI
                        result = main_MANDIRI(filepath, select, model)
                        show(result)
                        label(result[1], result[5]) 
                        look(result)                                                      
                    except: fail()

                elif sum(select_list)!=0 and max(select_list) > 2 and select == 104:  
                    try:
                        print("BNI2 file detected...")  
                        from BNI import main_BNI
                        result = main_BNI(filepath, select, model)
                        show(result)
                        label(result[1], result[5]) 
                        look(result)                             
                    except: fail()       
                    
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 105:  
                    try:
                        print("BNI3 file detected...") 
                        from BNI import main_BNI
                        result = main_BNI(filepath, select, model)
                        show(result)
                        label(result[1], result[5]) 
                        look(result)                             
                    except: fail() 

                elif sum(select_list)!=0 and max(select_list) > 2 and select == 106:  
                    try:
                        print("PERMATA2 file detected...")
                        print("Program is currently being developed...")
                        #from PERMATA import main_PERMATA
                        #result = main_PERMATA(filepath, select, model)                        
                        #SAMPAI SINI???
                    except: fail()                      
                    
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 107:  
                    try:
                        print("CIMB1 file detected...")
                        print("Program is currently being developed...")
                        #from CIMB import main_CIMB
                        #result = main_CIMB(filepath, select, model)                        
                        #SAMPAI SINI???
                    except: fail()
                    
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 108:  
                    try:
                        print("OCBC1 file detected...")
                        print("Program is currently being developed...")
                        #from OCBC import main_OCBC
                        #result = main_OCBC(filepath, select, model)                        
                        #SAMPAI SINI???
                    except: fail()  
                    
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 109:  
                    try:
                        print("DBS1 file detected...")
                        print("Program is currently being developed...")
                        #from DBS import main_DBS
                        #result = main_DBS(filepath, select, model)                        
                        #SAMPAI SINI???
                    except: fail()                  

                elif sum(select_list)!=0 and max(select_list) > 2 and select == 110:  
                    try:
                        print("QNB1 file detected...")
                        print("Program is currently being developed...")
                        #from QNB import main_QNB
                        #result = main_QNB(filepath, select, model)                        
                        #SAMPAI SINI???
                    except: fail()                     
                    
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 111:  
                    try:
                        print("UOB1 file detected...")
                        print("Program is currently being developed...")
                        #from UOB import main_UOB
                        #result = main_UOB(filepath, select, model)                        
                        #SAMPAI SINI???
                    except: fail() 
                    
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 112:  
                    try:
                        print("DANAMON1 file detected...")                   
                        from DANAMON import main_DANAMON
                        result = main_DANAMON(filepath, select, model)                        
                        show(result)
                        label(result[1], result[5]) 
                        look(result)                             
                    except: fail()                     
                    
                elif sum(select_list)!=0 and max(select_list) > 2 and select == 113:  
                    try:
                        print("HSBC1 file detected...")
                        print("Program is currently being developed...")
                        from HSBC import main_HSBC
                        result = main_HSBC(filepath, select, model)                        
                        #SAMPAI SINI???
                    except: fail() 
                        
                    
                else: fail_recognize() 
                
            except IndexError:
                print("")
                gui.text.insert("end", "unable to recognize pdf template..", "red", "\n") 
                print("") 
                gui.label3.config(text='extraction failed!',  foreground="red")            
    
    
    #==========gather info from 1st page ===========
        
        elif text != None or text != "":
            #print(text)
            found_BCA1 = len(re.findall(r"REKENING\s+TAHAPAN|Apabila\s+nasabah\s+tidak|BCA\s+berhak|akhir bulan berikutnya|Bersambung ke Halaman berikut|TANGGAL\s+KETERANGAN\s+CBG MUTASI\s+SALDO|nasabah\s+dianggap\s+telah\s+menyetujui\s+segala\s+data|Bersambung\s+ke\s+Halaman\s+berikut", text))
            found_BCA2 = len(re.findall(r"Informasi Rekening|Kode Mata Uang|Keterangan Cabang Jumlah Saldo|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))\s+(.*)\s+(\d{1,3},\d{1,3},\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3}\.\d\d|\d{1,3}\.\d\d)CR)", text))
            found_BRI1 = len(re.findall(r"LAPORAN TRANSAKSI BANK BRI| Dengan Setulus Hati|Tanggal Transaksi Uraian Transaksi Chq Debet Kredit Saldo Teller|(((0[1-9]|1\d|2\d|3[01])/(0?[1-9]|1[0-2])/\d{2})\s+(\d\d:\d\d:\d\d)\s+(.*)\s+(0.00)\s+(\d{1,3},\d{1,3},\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3}\.\d\d|\d{1,3}\.\d\d)K\s+)", text))
            found_BRI2 = len(re.findall(r"PT. BANK RAKYAT INDONESIA (PERSERO) TBK|DATE TIME REMARK DEBET CREDIT Ledger TELLER ID|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{2})\s+(\d\d:\d\d:\d\d)\s+(.*)\s+(0.00)\s+(\d{1,3},\d{1,3},\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3},\d{1,3}\.\d\d|\d{1,3},\d{1,3}\.\d\d|\d{1,3}\.\d\d)\s+)", text))
            found_MANDIRI1 = len(re.findall(r"TRANSACTION INQUIRY|Value Date Description Reference No. Debit Credit Saldo|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+(\d\d\.\d\d\.\d\d)\s+((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})(.*)\s+(0.00)\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d))", text))
            found_MANDIRI2 = len(re.findall(r"Rekening Koran|Reference No. Debit Credit Balance|((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})(.*)\s+(0.00)\s+((\d+,).*|\d+\.\d\d)", text))
            found_MANDIRI3 = len(re.findall(r"Laporan Rekening Koran|Account Statement Report|Posting Date Remark Reference No Debit Credit Balance|((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+\d+:.*\n(.*)(0.00)\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)|(\d+\s+\w+\s+\d{4})\s+(.*)\s+(0.00)\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)|(\d+\s+\w+\s+\d{4})\s+(.*)\s+((\d+,).*|\d+\.\d\d)\s+(0.00)\s+((\d+,).*|\d+\.\d\d)", text))
            found_BNI1 = len(re.findall(r"No. Post Date Branch Journal No.|\d+\s+((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+(\w+)\s+(\d{6})\s+(.*)\s+((\d+,).*|\d+\.\d\d)\s+(C|D)\s+((\d+,).*|\d+\.\d\d)", text))
            found_BNI2 = len(re.findall(r"Posting Date Effective Date|((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+\d{2}\.\d{2}\.\d{2}\s+((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2])/\d{4})\s+\d{2}\.\d{2}\.\d{2}\s+", text))
            found_PERMATA1 = len(re.findall(r"Transaction\s+Cheque\s+Number\s+Customer|Post\s+Date\s+Eff\s+Date\s+Ref\s+No\s+Description\s+Debit\s+Credit|\d+((0[1-9]|1\d|2\d|3[01])-\w{3}-\d{4})\s+((0[1-9]|1\d|2\d|3[01])-\w{3}-\d{4})\s+", text))
            found_CIMB1 = len(re.findall(r"Laporan\s+Konsolidasi|Consolidated\s+Account\s+Statement|No.\s+Rekening\s+Nama\s+Product\s+Mata\s+Uang\s+Saldo|Account\s+Number\s+Product\s+Name\s+Currency\s+Balance|Average\s+Balance|Saldo\s+Rata-Rata|Laporan\s+transaksi\s+ini\s+akan\s+dikirim\s+ke\s+alamat|Mohon\s+periksa\s+laporan\s+transaksi\s+ini|Call\s+Centre\s+di\s+nomor\s+telepon\s+14041|authorised\s+signatory\s+of\s+PT.\s+Bank\s+CIMB\s+Niaga\s+Tbk|Tgl.\s+Txn\s+Tgl.\s+Valuta\s+Uraian\s+Transaksi|Txn.\s+Date\s+Val.\s+Date\s+Description", text))
            found_CIMB2 = len(re.findall(r"No.\s+Post\s+Date\s+Eff\s+Date\s+Cheque\s+No\s+Description|Description\s+Debit\s+Credit\s+Balance\s+Transaction\s+Ref\s+No|\d+\s+(\d{2}/(0[1-9]|1\d|2\d|3[01])/\d{2})\s+\d{2}:\d{2}\s+(\d{2}/(0[1-9]|1\d|2\d|3[01])/\d{2})\s+\d{2}:\d{2}\s+(.*)\s+0.00\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)|\d+\s+(\d{2}/(0[1-9]|1\d|2\d|3[01])/\d{2})\s+\d{2}:\d{2}\s+(\d{2}/(0[1-9]|1\d|2\d|3[01])/\d{2})\s+\d{2}:\d{2}\s+(.*)\s+((\d+,).*|\d+\.\d\d)\s+0.00\s+((\d+,).*|\d+\.\d\d)", text))
            found_OCBC1 = len(re.findall(r"Account No :.*Opening Balance :.*|Account Name.*Closing Balance :|Total Debit.*Ledger Balance :|Total Credit :.*Available Balance :|Value Date Reference No. Cheque No. Description Debit Credit Balance|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+\d{10,}\s+.*\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)", text))
            found_OCBC2 = len(re.findall(r"From: \d\d-\w+-\d{4} To: \d\d-\w+-\d{4}|Post Date Value Date Ref/Cheque No. Description Debit Credit Balance|\d\d-\w+-\d{4} \d\d-\w+-\d{4}\s+\d{16}\s+.*\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)", text))
            found_DBS1 = len(re.findall(r"Account Number :.*Account Name :|Product Type :|Opening Balance.*Earmark Amount|Ledger Balance.*Overdraft Limit|Date Value Date Transaction Details  Debit Credit Running Balance", text)) + len(re.findall(r"Transactions performed on a non-working day|please select the next business day to view your transaction", text_end))
            found_QNB1 = len(re.findall(r"No. Post Date Effective Date Transaction Name Debit Credit Balance Information|layanan CIB BANK QNB INDONESIA|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+\d{2,5}\s+.*\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)\s+\w+", text))
            found_UOB1 = len(re.findall(r"CompanyID:.*AccountNumber:|AccountName:.*LedgerBalance:|AccountType:.*AvailableBalance:|AccountCurrency:.*TotalFloat:|AccountBranch:.*OverdraftFacility:|AccountNature.*Earmark|Primary.*AllocatedAmount|Movement Details - From:|Statement\s+ValueDate\s+Transaction\s+Description\s+Deposit\s+Withdrawal\s+Balance|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(.*)\s+(.*)\s+0\s+((\d+,).*|\d+)\s+(-(\d+,).*|-\d+|(\d+,).*|\d+)|(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(((0[1-9]|1\d|2\d|3[01])/(0[1-9]|1[0-2]))/\d{4})\s+(.*)\s+(.*)\s+((\d+,).*|\d+)\s+0\s+(-(\d+,).*|-\d+|(\d+,).*|\d+)", text))
            found_DANAMON1 = len(re.findall(r"Transaction Inquiry Report|Account Number :(.*)Period : \w+ \d\d, \d{4} - \w+ \d\d, \d{4}|Name :(.*)Print Date : \w+ \d\d, \d{4} \d\d:\d\d|Posting Date Value Date Transaction Branch Reference Number Description Debit Credit Balance|(0[1-9]|1\d|2\d|3[01])\s+\w+\s+\d{4}\s+(0[1-9]|1\d|2\d|3[01])\s+\w+\s+\d{4}\s+(.*)\s+0.00\s+((\d+,).*|\d+\.\d\d)\s+((\d+,).*|\d+\.\d\d)", text))
            found_HSBC1 = len(re.findall(r"Laporan Rekening$|Account Statement$|Rincian Rekening$|Account Details$|Nomor Urut Laporan|Stmt Sequence Number|Tanggal Rincian Penarikan Setoran Saldo|Date Details Withdrawals Deposits Balance|(0[1-9]|1\d|2\d|3[01])\w{3}\d{4}\s+(.*)\s+(-(\d+,).*|(\d+,).*|-\d+|\d+)|Bank HSBC Indonesia|Kantor Pusat - World Trade Center|HSBC Premier|HSBC Advance|Layanan Nasabah\s*:\s*www.hsbc.co.id|Diterbitkan oleh PT Bank HSBC Indonesia|6221 2551 4722|6221 2552 6603|1 500 700|1 500 808", text))
            
            select_list = [found_BCA1, found_BCA2, found_BRI1, found_BRI2, found_MANDIRI1, found_MANDIRI2, found_MANDIRI3, found_BNI1, found_BNI2, found_PERMATA1, found_CIMB1, found_CIMB2, found_OCBC1, found_OCBC2, found_DBS1, found_QNB1, found_UOB1, found_DANAMON1, found_HSBC1]
            select = select_list.index(max(select_list))
            print(select_list)
            
            #==================pdfplumber==========================
            if sum(select_list)!=0 and max(select_list) > 2 and select == 0: #BCA1
                try:
                    print("BCA1 file detected...")
                    print("processing pdfplumber...")
                    from BCA import main_BCA
                    result = main_BCA(filepath, select, None)
                    show(result)
                    label(result[1], result[5])   
                    look(result)          
                except: fail() 
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==1: #BCA2
                try:
                    print("BCA2 file detected...")
                    print("processing pdfplumber...")
                    from BCA import main_BCA
                    result = main_BCA(filepath, select, None)
                    show(result)
                    label(result[1], result[5])   
                    look(result)                          
                except: fail()         
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==2: #BRI1
                try:                
                    print("BRI1 file detected...")
                    print("processing pdfplumber...")
                    from BRI import main_BRI              
                    result = main_BRI(filepath, select, None)
                    show(result)   
                    label(result[1], result[5])   
                    look(result)                           
                except: fail()     
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==3: #BRI2   
                try:                
                    print("BRI2 file detected...")
                    print("processing pdfplumber...")
                    from BRI import main_BRI
                    result = main_BRI(filepath, select, None)
                    show(result)                      
                    label(result[1], result[5])   
                    look(result)                           
                except: fail()          
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==4: #MANDIRI1
                try:                 
                    print("MANDIRI1 file detected...")
                    print("processing pdfplumber...")   
                    from MANDIRI import main_MANDIRI
                    result = main_MANDIRI(filepath, select, None)                
                    show(result)                    
                    label(result[1], result[5])   
                    look(result)                           
                except: fail()                 
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==5: #MANDIRI2
                try:                 
                    print("MANDIRI2 file detected...")            
                    print("processing pdfplumber...") 
                    from MANDIRI import main_MANDIRI
                    result = main_MANDIRI(filepath, select, None)                
                    show(result)
                    label(result[1], result[5])   
                    look(result)                          
                except: fail()                 
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==6: #MANDIRI3
                try:                   
                    print("MANDIRI3 file detected...")            
                    print("processing pdfplumber...")                 
                    from MANDIRI import main_MANDIRI
                    result = main_MANDIRI(filepath, select, None)                
                    show(result)
                    label(result[1], result[5])   
                    look(result)                           
                except: fail()                    
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==7:#BNI1
                try:                 
                    print("BNI1 file detected...")            
                    print("processing pdfplumber...")                                     
                    from BNI import main_BNI  
                    result = main_BNI(filepath, select, None)     
                    show(result)
                    label(result[1], result[5])   
                    look(result)                          
                except: fail() 
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==8:#BNI2
                try:
                    print("BNI2 file detected...")            
                    print("processing pdfplumber...")                       
                    from BNI import main_BNI   
                    result = main_BNI(filepath, select, None)    
                    show(result)
                    label(result[1], result[5])   
                    look(result)                          
                except: fail()                
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==9:#PERMATA1
                try:
                    print("PERMATA1 file detected...")            
                    print("processing pdfplumber...")                   
                    from PERMATA import main_PERMATA   
                    result = main_PERMATA(filepath, select, None)                 
                    show(result)
                    label(result[1], result[6])         
                    if not re.findall('#  WARNING', result[1]):
                        gui.link.bind("<Button-1>", lambda event, : os.startfile(result[6], 'open') )                
                        gui.link.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))                 
                        gui.link2.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
                        gui.link2.config(text='open file input', fg="green", cursor="hand2", font=('Calibri',11,'underline'))                          
                except: fail()    
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==10: #CIMB1
                try:
                    print("CIMB1 file detected...")            
                    print("processing pdfplumber...")                  
                    from CIMB import main_CIMB   
                    result = main_CIMB(filepath, select, None)                  
                    show(result)
                    label(result[1], result[5])                       
                    look(result)                           
                except: fail()                   
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==11: #CIMB2
                try:
                    print("CIMB2 file detected...")            
                    print("processing pdfplumber...") 
                    from CIMB import main_CIMB   
                    result = main_CIMB(filepath, select, None)                 
                    show(result)
                    label(result[1], result[5])                      
                    look(result)                           
                except: fail()                 
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==12:#OCBC1
                try:
                    print("OCBC1 file detected...")            
                    print("processing pdfplumber...") 
                    from OCBC import main_OCBC   
                    result = main_OCBC(filepath, select, None)                  
                    show(result)
                    label(result[1], result[5])                      
                    look(result)                  
                except: fail()  
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==13: #OCBC2
                try:
                    print("OCBC2 file detected...")            
                    print("processing pdfplumber...") 
                    from OCBC import main_OCBC   
                    result = main_OCBC(filepath, select, None)  
                    show(result)
                    label(result[1], result[5])                       
                    look(result)                
                except: fail()                
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==14:#DBS1
                try:
                    print("DBS1 file detected...")            
                    print("processing pdfplumber...")  
                    from DBS import main_DBS
                    result = main_DBS(filepath, select, None)  
                    show(result)
                    label(result[1], result[5])                       
                    look(result)                  
                except: fail()             
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==15 and not re.findall(': USD|: EUR|: GBP|(USD)|(EUR)|(GBP)', text):#QNB1
                try:
                    print("QNB1 file detected...")            
                    print("processing pdfplumber...")  
                    from QNB import main_QNB
                    result = main_QNB(filepath, select, None)  
                    show(result)
                    label(result[1], result[5])                       
                    look(result)                 
                except: fail()                   
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==16: #UOB1
                try:
                    print("UOB1 file detected...")            
                    print("processing pdfplumber...")  
                    from UOB import main_UOB
                    result = main_UOB(filepath, select, None)                  
                    show(result)
                    label(result[1], result[5])                       
                    look(result)                 
                except: fail()                 
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==17: #DANAMON1
                try:
                    print("DANAMON1 file detected...")            
                    print("processing pdfplumber...")
                    from DANAMON import main_DANAMON
                    result = main_DANAMON(filepath, select, None)
                    show(result)
                    label(result[1], result[5])                        
                    look(result)                 
                except: fail()   
            elif sum(select_list)!=0 and max(select_list) > 2 and select ==18: #HSBC1
                try:
                    print("HSBC1 file detected...")            
                    print("processing pdfplumber...")
                    from HSBC import main_HSBC
                    result = main_HSBC(filepath, select, None)
                    show(result)
                    label(result[1], result[5])                       
                    look(result)                     
                except: fail()                     
                    
                    
            #======================none of above======================
            else: fail_recognize()      
    
    except FileNotFoundError:
        gui.text.insert("end", "No file selected", "red", "\n")
        print("")
        gui.label3.config(text='process aborted!',  foreground="red")
        pass                   
    else:
        pass     
    
    time.sleep(0.) 

def func_b():   
    
    def show(result):
        
        if re.findall('No ', result):
            gui.text.insert("end", result, "green")
        elif re.findall('Found', result):
            gui.text.insert("end", result, "blue")
        elif re.findall('Unable', result):
            gui.text.insert("end", result, "red") 
            gui.label3.config(text='process aborted!',  foreground="red")            
    
    #=================================================start/body================================================
    #clear all status
    gui.label1.config(text='')
    gui.label3.config(text='')
    gui.link.config(text='')
    gui.link1.config(text='')
    gui.link2.config(text='')
    
    filepath= fd.askopenfilename(title='Open excel file', filetypes=(('excel files', '*.xlsx'), ('excel files', '*.xls')))
    onlyfilename = os.path.basename(filepath) 
    onlypath = os.path.dirname(filepath) #filepath=onlypath + "/" + onlyfilename
    gui.label1.config(text='file selected : '+ onlyfilename, foreground="blue")
    print("")
    print("source: ",filepath)
    #start progbar
    gui.progbar.start() 
    try:
        try:
            os.rename(filepath,filepath)            
            from uturn import uturn
            result = uturn(filepath)
            show(result)   
            if not re.findall('Unable', result):            
                gui.link1.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
                gui.link1.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))       
    
        except PermissionError:
            gui.text.insert("end", "please close excel file output", "red", "\n")
            gui.label3.config(text='process aborted!',  foreground="red")
            print("")
            sys.exit(1)            
    except FileNotFoundError:
        gui.text.insert("end", "No file selected", "red", "\n") 
        print("")
        gui.label3.config(text='process aborted!',  foreground="red")
        pass                   
    else:
        pass     
    time.sleep(0.) 
    
    
def func_c(relpar_names):   

    def show(result):
        
        if re.findall('No ', result):
            gui.text.insert("end", result, "green")
        elif re.findall('Found', result):
            gui.text.insert("end", result, "blue")
        elif re.findall('Unable', result):
            gui.text.insert("end", result, "red") 
            gui.label3.config(text='process aborted!',  foreground="red")
            
    #=================================================start/body================================================
    #clear all status
    gui.label1.config(text='')
    gui.label3.config(text='')
    gui.link.config(text='')
    gui.link1.config(text='')
    gui.link2.config(text='')    
    
    filepath= fd.askopenfilename(title='Open excel file', filetypes=(('excel files', '*.xlsx'), ('excel files', '*.xls')))
    onlyfilename = os.path.basename(filepath) 
    onlypath = os.path.dirname(filepath) #filepath=onlypath + "/" + onlyfilename
    gui.label1.config(text='file selected : '+ onlyfilename, foreground="blue")
    print("")
    print("source: ",filepath)
    #start progbar
    gui.progbar.start()
    print("searching related party names: ", relpar_names)
    
    try:
        try:
            os.rename(filepath,filepath)            
            from relpar import relpar
            result = relpar(filepath, relpar_names)
            show(result)   
            if not re.findall('Unable', result):
                gui.link1.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
                gui.link1.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))  
        except PermissionError:
            gui.text.insert("end", "please close excel file output", "red", "\n")
            gui.label3.config(text='process aborted!',  foreground="red")
            print("")
            sys.exit(1)            
    except FileNotFoundError:
        gui.text.insert("end", "No file selected", "red", "\n") 
        print("")
        gui.label3.config(text='process aborted!',  foreground="red")
        pass                   
    else:
        pass        
    time.sleep(0.)     
  
    
def func_d(majbuy_names):   

    def show(result):
        
        if re.findall('No ', result):
            gui.text.insert("end", result, "green")
        elif re.findall('Found', result):
            gui.text.insert("end", result, "blue")
        elif re.findall('Unable', result):
            gui.text.insert("end", result, "red") 
            gui.label3.config(text='process aborted!',  foreground="red")
            
    #=================================================start/body================================================
    #clear all status
    gui.label1.config(text='')
    gui.label3.config(text='')
    gui.link.config(text='')
    gui.link1.config(text='')
    gui.link2.config(text='')    
    
    filepath= fd.askopenfilename(title='Open excel file', filetypes=(('excel files', '*.xlsx'), ('excel files', '*.xls')))
    onlyfilename = os.path.basename(filepath) 
    onlypath = os.path.dirname(filepath) #filepath=onlypath + "/" + onlyfilename
    gui.label1.config(text='file selected : '+ onlyfilename, foreground="blue")
    print("")
    print("source: ",filepath)
    #start progbar
    gui.progbar.start()
    
    print("searching related party names: ", majbuy_names)
    
    try:
        try:
            os.rename(filepath,filepath)             
            from majbuy import majbuy
            result = majbuy(filepath, majbuy_names)
            show(result)   
            if not re.findall('Unable', result):            
                gui.link1.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
                gui.link1.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))  
            
        except PermissionError:
            gui.text.insert("end", "please close excel file output", "red", "\n")
            gui.label3.config(text='process aborted!',  foreground="red")
            print("")
            sys.exit(1)            
    
    except FileNotFoundError:
        gui.text.insert("end", "No file selected", "red", "\n") 
        print("")
        gui.label3.config(text='process aborted!',  foreground="red")
        pass                   
    else:
        pass        
    
    time.sleep(0.)
    
def func_e(majsup_names):   

    def show(result):
        
        if re.findall('No ', result):
            gui.text.insert("end", result, "green")
        elif re.findall('Found', result):
            gui.text.insert("end", result, "blue")
        elif re.findall('Unable', result):
            gui.text.insert("end", result, "red") 
            gui.label3.config(text='process aborted!',  foreground="red")
            
    #=================================================start/body================================================
    #clear all status
    gui.label1.config(text='')
    gui.label3.config(text='')
    gui.link.config(text='')
    gui.link1.config(text='')
    gui.link2.config(text='')    
    
    filepath= fd.askopenfilename(title='Open excel file', filetypes=(('excel files', '*.xlsx'), ('excel files', '*.xls')))
    onlyfilename = os.path.basename(filepath) 
    onlypath = os.path.dirname(filepath) #filepath=onlypath + "/" + onlyfilename
    gui.label1.config(text='file selected : '+ onlyfilename, foreground="blue")
    print("")
    print("source: ",filepath)
    
    #start progbar
    gui.progbar.start()
    
    print("searching related party names: ", majsup_names)
    
    try:
        try:
            os.rename(filepath,filepath)             
            from majsup import majsup
            result = majsup(filepath, majsup_names)
            show(result)  
            if not re.findall('Unable', result):            
                gui.link1.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
                gui.link1.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))  
            
        except PermissionError:
            gui.text.insert("end", "please close excel file output", "red", "\n")
            gui.label3.config(text='process aborted!',  foreground="red")
            print("")
            sys.exit(1)            
    
    except FileNotFoundError:
        gui.text.insert("end", "No file selected", "red", "\n") 
        print("")
        gui.label3.config(text='process aborted!',  foreground="red")
        pass                   
    else:
        pass        
    
    time.sleep(0.)

    
def func_f(any_names):   

    def show(result):
        
        if re.findall('No ', result):
            gui.text.insert("end", result, "green")
        elif re.findall('Found', result):
            gui.text.insert("end", result, "blue")
        elif re.findall('Unable', result):
            gui.text.insert("end", result, "red") 
            gui.label3.config(text='process aborted!',  foreground="red")
            
    #=================================================start/body================================================
    #clear all status
    gui.label1.config(text='')
    gui.label3.config(text='')
    gui.link.config(text='')
    gui.link1.config(text='')
    gui.link2.config(text='')    
    
    filepath= fd.askopenfilename(title='Open excel file', filetypes=(('excel files', '*.xlsx'), ('excel files', '*.xls')))
    onlyfilename = os.path.basename(filepath) 
    onlypath = os.path.dirname(filepath) #filepath=onlypath + "/" + onlyfilename
    gui.label1.config(text='file selected : '+ onlyfilename, foreground="blue")
    print("")
    print("source: ",filepath)
    #start progbar
    gui.progbar.start()    
    print("searching related party names: ", any_names)
    
    try:
        try:         
            os.rename(filepath,filepath)              
            from anynames import anynames
            result = anynames(filepath, any_names)
            show(result)   
            if not re.findall('Unable', result):            
                gui.link1.bind("<Button-1>", lambda event, : os.startfile(filepath, 'open') )                
                gui.link1.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))  
            
        except PermissionError:
            gui.text.insert("end", "please close excel file output", "red", "\n")
            gui.label3.config(text='process aborted!',  foreground="red")
            print("")
            sys.exit(1)            
    
    except FileNotFoundError:
        gui.text.insert("end", "No file selected", "red", "\n") 
        print("")
        gui.label3.config(text='process aborted!',  foreground="red")
        pass                   
    else:
        pass        
    time.sleep(0.)
        
def func_g(keywords): 

    document = Document()
    #start progbar
    gui.progbar.start()       
    query = keywords[0]

    try:
        for j in search(query, tld="co.id", num=20, stop=20, pause=4):
            try:
                gui.text.insert("end", j, "green", "\n") 
                p = document.add_paragraph('source: ')
                p.add_run(j).underline = True   
                p = document.add_paragraph('\n')
                url = j
                
                #if using plain article
                #article = Article(url)
                #article.download()
                #article.html
                #article.parse()
                #content = article.text
                
                #if using beautifulsoup4
                #r = requests.get(url)
                #soup = BeautifulSoup(r.text, 'html5lib')
                #content = soup.prettify()

                # if using newspaper3k
                #headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 6.1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/41.0.2228.0 Safari/537.3'}
                #req = Request(url=url, headers=headers) 
                #page = urlopen(req)
                #page_content = page.read()
                
                #with open('page_content.html', 'w') as fid:
                #    fid.write(str(page_content))   
                
                #article = Article('')
                #article.set_html(open("page_content.html").read())
                #article.parse()
                #content = article.text

                #print(content)
                #p.add_run(content)
            except:
                gui.text.insert("end", "unable to load link..", "red", "\n")   
                pass

        try:
            files = [('Word document', '*.docx'), ]
            filename = asksaveasfile(filetypes = files, defaultextension = files)
            name=str(filename.name)
            document.save(name)
            #show link to output
            gui.link1.bind("<Button-1>", lambda event, : os.startfile(name, 'open') )                
            gui.link1.config(text='open file output', fg="green", cursor="hand2", font=('Calibri',11,'underline'))     

        except PermissionError:
            gui.text.insert("end", "please close Word Document", "red", "\n")
        except AttributeError:
            gui.text.insert("end", "No file save", "red", "\n")
    except:
        gui.text.insert("end", "unable to connect internet, please check your connection..", "red", "\n")
         
    time.sleep(0.)
    
#=====================================
def arbitrarySLIK():
    func_aSLIK()
def arbitrary():
    func_a()
def arbitrary2():
    func_b()
def arbitrary3(relpar_names):
    func_c(relpar_names)    
def arbitrary4(majbuy_names):
    func_d(majbuy_names)
def arbitrary5(majsup_names):
    func_e(majsup_names)    
def arbitrary6(any_names):
    func_f(any_names)     
def arbitrary7(keywords):
    func_g(keywords)     

gui = GUI_Core()
gui.root.mainloop()   

sys.stdout = old_stdout 


# In[ ]:




